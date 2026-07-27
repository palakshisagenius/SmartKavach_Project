# generate_report.py
# Generates the SmartKavach Final Technical Report as a Word document

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Helper functions ─────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1C, 0x4E, 0x80) if level==1 else RGBColor(0x2E, 0x75, 0xB6)
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if header:
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml.ns import qn
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1C4E80')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
    return row

def sp():
    doc.add_paragraph()

# ── COVER PAGE ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SMARTKAVACH")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1C, 0x4E, 0x80)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("AI-Enhanced Train Collision Avoidance System")
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Final Technical Report")
run3.font.size = Pt(16)
run3.italic = True

sp()
t = doc.add_table(rows=0, cols=2)
t.style = 'Table Grid'
for label, value in [
    ("Student", "Palaksh Bhardwaj"),
    ("Programme", "B.Tech CSE (Artificial Intelligence)"),
    ("Institution", "Gautam Budh University, Greater Noida"),
    ("Academic Year", "2025–2026"),
    ("Date", "July 2026"),
]:
    row = t.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
heading("Abstract")
para("SmartKavach is an AI-enhanced overlay on India's KAVACH / TCAS (Train Collision Avoidance System) that introduces predictive, adaptive, and generative capabilities to the existing rule-based railway safety platform. The system augments KAVACH with three machine learning models: an XGBoost regressor for dynamic Movement Authority prediction (RMSE 19.05m, R² 0.99), an LSTM Autoencoder for real-time anomaly detection (Recall 1.00, Accuracy 95%), and a Random Forest regressor for context-aware adaptive speed profiling (RMSE 0.72 km/h, R² 0.9986). All models were trained on a 10,000-event synthetic dataset modelled on the KAVACH data architecture and validated against real RDSO trial data from the NCR Division (July 2024), achieving less than 10% error on real observed Movement Authority values. The system is deployed as a FastAPI inference server with three REST endpoints, a live React NMS dashboard, and a real-time data simulator. A suite of 10 automated tests confirms all system performance targets are met.")

doc.add_page_break()

# ── 1. INTRODUCTION ──────────────────────────────────────────────────────────
heading("1. Introduction")
heading("1.1 Background", 2)
para("India's KAVACH system, developed by RDSO (Research Designs and Standards Organisation), is the national Automatic Train Protection platform. It prevents Signal Passed At Danger (SPAD) events, enforces speed limits, and prevents head-on and rear-end collisions through radio-based train position awareness and automatic brake application. The system comprises Onboard KAVACH units (Loco TCAS), Stationary KAVACH units (STCAS), RFID positioning tags, UHF radio communication towers, and a centralised Network Management System (NMS).")

heading("1.2 Problem Statement", 2)
para("While KAVACH successfully enforces fixed safety rules, it operates as a deterministic, reactive system with four key limitations:")
bullet("Movement Authority is computed solely from signal relay data, without contextual factors such as weather, freight load, or track condition history")
bullet("The system reacts to danger thresholds rather than predicting emerging risk scenarios before they escalate")
bullet("The NMS logs and monitors events but cannot predict faults, identify recurring anomaly patterns, or perform root-cause analysis")
bullet("Speed supervision uses static section limits that do not adapt to real-world operational conditions")

heading("1.3 Objectives", 2)
para("SmartKavach addresses these limitations by adding an AI intelligence layer that:")
bullet("Computes dynamic Movement Authority using ML models trained on contextual features")
bullet("Detects anomalies in real time using LSTM Autoencoders trained on normal operational patterns")
bullet("Generates adaptive speed profiles sensitive to weather, load, and section incident history")
bullet("Provides an intelligent NMS dashboard with predictive fault detection")

doc.add_page_break()

# ── 2. LITERATURE REVIEW ─────────────────────────────────────────────────────
heading("2. Literature Review and Related Work")
para("Automatic Train Protection systems have evolved significantly from fixed-block signalling to communication-based moving block systems. European Train Control System (ETCS) and Japan's Digital ATC represent the state of the art in deterministic ATP. However, the integration of machine learning into railway safety systems remains an emerging area.")
para("Recent work in railway anomaly detection has applied LSTM networks to track circuit failure prediction and train delay forecasting. XGBoost has been applied to train energy consumption optimisation and delay prediction with strong results on tabular railway data. Random Forest models have been used for track degradation prediction in maintenance scheduling contexts.")
para("SmartKavach distinguishes itself by targeting the specific KAVACH architecture — using RFID positioning data, radio packet streams, and signal relay information as input features — and by validating synthetic model training against real RDSO trial data, a contribution not present in existing literature on Indian railway AI applications.")

doc.add_page_break()

# ── 3. SYSTEM DESIGN ─────────────────────────────────────────────────────────
heading("3. System Design and Architecture")
heading("3.1 Design Principles", 2)
bullet("Non-invasive: SmartKavach reads from existing data streams and produces advisory outputs only. It never overrides mandatory KAVACH hardware commands")
bullet("Fail-safe: Any unavailability of the AI layer causes transparent fallback to standard KAVACH operation within one second")
bullet("Modular: Each AI component is independently deployable and testable")
bullet("Explainable: All predictions include confidence scores for operator review")

heading("3.2 Five-Layer Architecture", 2)
t2 = doc.add_table(rows=0, cols=3)
t2.style = 'Table Grid'
table_row(t2, ["Layer", "Components", "Role"], header=True)
for row_data in [
    ["1 — Physical", "RFID tags, Onboard KAVACH, Stationary KAVACH, Loco sensors", "Data generation (existing KAVACH hardware)"],
    ["2 — Communication", "GSM/LTE, GPS/GNSS, OFC/E1, Key Management Server", "Data transport (existing infrastructure)"],
    ["3 — AI Intelligence", "Predictive MA Engine, Anomaly Detector, Speed Profiler", "AI processing (new SmartKavach layer)"],
    ["4 — NMS & Monitoring", "NMS server, AI Dashboard, Historical Data Store", "Centralised monitoring (enhanced)"],
    ["5 — External Feeds", "Weather API, track condition DB, incident history", "Contextual enrichment (new data sources)"],
]:
    table_row(t2, row_data)

heading("3.3 Technology Stack", 2)
t3 = doc.add_table(rows=0, cols=2)
t3.style = 'Table Grid'
table_row(t3, ["Component", "Technology"], header=True)
for row_data in [
    ["MA Prediction Model", "XGBoost (gradient boosting regression)"],
    ["Anomaly Detection", "LSTM Autoencoder (TensorFlow/Keras)"],
    ["Speed Profiler", "Random Forest Regressor (scikit-learn)"],
    ["API Server", "FastAPI with Pydantic schemas"],
    ["Dashboard Frontend", "React with Recharts and Leaflet"],
    ["Data Pipeline", "Python simulator, REST/JSON"],
]:
    table_row(t3, row_data)

doc.add_page_break()

# ── 4. IMPLEMENTATION ────────────────────────────────────────────────────────
heading("4. Implementation")
heading("4.1 Synthetic Dataset Generation", 2)
para("Since real KAVACH operational logs are not publicly available, a Python simulator was developed to generate a 10,000-event synthetic dataset modelling the KAVACH data flow described in RDSO technical documentation. The dataset spans 10 trains across 5 track sections, with 5% injected anomalies of four types: RFID read failure, radio communication drop, overspeed, and sudden brake application. Each event contains 16 features including speed, signal aspect, RFID read success rate, radio packet loss, weather index, freight load, deceleration rate, and labelled targets for all three models.")
para("All four anomaly types were subsequently confirmed present in real KAVACH RDSO trial data (NCR Division, July 2024), validating the realism of the synthetic dataset design.")

heading("4.2 Movement Authority Prediction Model", 2)
para("An XGBoost regressor was trained on 8,000 events to predict safe Movement Authority distance in metres. Input features include speed, signal aspect (label encoded), weather index, freight load, section speed limit, section incident rate, RFID read success, radio packet loss, and deceleration rate. The model achieved RMSE of 19.05 metres and R² of 0.99 on the held-out test set. Feature importance analysis confirmed that speed, signal aspect, and weather index are the three most influential predictors, consistent with the physics of train braking distance.")

heading("4.3 Anomaly Detection Model", 2)
para("An LSTM Autoencoder was trained exclusively on normal operation sequences. The model takes sliding windows of 30 consecutive events across 6 features and learns to compress and reconstruct normal patterns. Anomalies produce high reconstruction error because the model was never trained on abnormal data. A threshold at the 95th percentile of training reconstruction errors (0.1790) separates normal from anomalous sequences. The model achieved recall of 1.00 — detecting every injected anomaly — with overall accuracy of 95%. Anomaly reconstruction errors were observed to be over 1,000 times higher than normal errors, indicating a clear and robust separation.")

heading("4.4 Adaptive Speed Profiler", 2)
para("A Random Forest regressor with 100 estimators was trained to predict contextually safe advisory speeds per train-section combination. A hard cap ensures advisory speed never exceeds the section speed limit, resulting in zero speed limit violations across all test cases. The model achieved RMSE of 0.72 km/h and R² of 0.9986. Feature importance confirmed speed, weather index, and freight load as the top three predictors.")

heading("4.5 FastAPI Inference Server", 2)
para("All three models are wrapped in a FastAPI server with three REST endpoints: POST /predict/ma, POST /detect/anomaly, and POST /predict/speed. Models are loaded once at server startup and kept in memory for low-latency inference. A sliding window buffer per train ID maintains the last 31 events for LSTM sequence construction. Pydantic schemas enforce strict request validation, returning HTTP 422 for malformed inputs and HTTP 500 with fallback logging for model failures.")

heading("4.6 React NMS Dashboard", 2)
para("A React dashboard visualises live predictions from all three models. Three train cards display real-time speed, advisory speed, Movement Authority, and weather severity. A Recharts line chart shows MA history for all trains over the last 20 update cycles. An anomaly alert feed displays the 10 most recent alerts with severity classification and timestamp. The dashboard polls the API every 4 seconds using Promise.all for parallel endpoint calls.")

doc.add_page_break()

# ── 5. RESULTS ───────────────────────────────────────────────────────────────
heading("5. Results and Evaluation")
heading("5.1 Model Performance", 2)
t4 = doc.add_table(rows=0, cols=4)
t4.style = 'Table Grid'
table_row(t4, ["Model", "Key Metric", "Value", "Target"], header=True)
for row_data in [
    ["MA Prediction (XGBoost)", "RMSE", "19.05 metres", "< 50m ✅"],
    ["MA Prediction (XGBoost)", "R²", "0.9900", "— ✅"],
    ["Anomaly Detection (LSTM)", "Recall", "1.00", "> 0.90 ✅"],
    ["Anomaly Detection (LSTM)", "Accuracy", "95%", "— ✅"],
    ["Speed Profiler (RF)", "RMSE", "0.72 km/h", "— ✅"],
    ["Speed Profiler (RF)", "R²", "0.9986", "— ✅"],
    ["Speed Profiler (RF)", "Violations", "0", "0 ✅"],
]:
    table_row(t4, row_data)

sp()
heading("5.2 Real-World Validation", 2)
para("The MA prediction model was evaluated against three real braking events from RDSO KAVACH trials conducted on the NCR Division in July 2024 (Loco 23545, Vande Bharat, RDE–VRBD section):")
t5 = doc.add_table(rows=0, cols=5)
t5.style = 'Table Grid'
table_row(t5, ["Speed (km/h)", "Real MA (m)", "Predicted MA (m)", "Error (m)", "Error %"], header=True)
for row_data in [
    ["115", "768", "795.2", "+27.2", "3.5%"],
    ["157", "1326", "1393.6", "+67.6", "5.1%"],
    ["158", "1516", "1393.6", "-122.4", "8.1%"],
]:
    table_row(t5, row_data)
sp()
para("Average absolute error: 72.4 metres across a real MA range of 768–1516 metres. All predictions are within 10% of real observed values, demonstrating strong generalisation from synthetic training data to real-world conditions.")

heading("5.3 System Performance", 2)
t6 = doc.add_table(rows=0, cols=3)
t6.style = 'Table Grid'
table_row(t6, ["Metric", "Target", "Result"], header=True)
for row_data in [
    ["MA endpoint latency", "< 500ms", "✅ Met"],
    ["Anomaly alert latency", "< 2 seconds", "✅ Met"],
    ["Speed limit violations", "0", "✅ 0 violations"],
    ["AI fallback time", "< 1 second", "✅ Met"],
    ["Automated test suite", "All pass", "✅ 10/10 passing"],
    ["Dashboard refresh", "Real-time", "✅ 4-second cycle"],
]:
    table_row(t6, row_data)

doc.add_page_break()

# ── 6. CONCLUSION ────────────────────────────────────────────────────────────
heading("6. Conclusion")
para("SmartKavach demonstrates that an AI intelligence layer can be successfully integrated with India's KAVACH train collision avoidance architecture without modifying any existing safety-certified hardware or logic. The system addresses four identified limitations of the current deterministic KAVACH design: static Movement Authority computation, reactive-only safety responses, limited NMS intelligence, and fixed speed profiles.")
para("All three machine learning models exceed their performance targets, and the MA prediction model achieves less than 10% error when validated against real RDSO trial data from July 2024 — a result that supports the practical viability of the approach. The complete system, from data pipeline through AI inference to live dashboard, is implemented and operational.")
para("Future work would focus on integration with real KAVACH NMS data streams, formal safety certification per Indian railway standards, and expansion to additional anomaly detection categories based on extended trial data analysis.")

heading("6.1 Key Contributions", 2)
bullet("First integration of LSTM Autoencoder anomaly detection with the KAVACH/TCAS data architecture")
bullet("Synthetic dataset validated against real RDSO trial data (NCR Division, July 2024)")
bullet("Complete working system: data pipeline, three AI models, REST API, live dashboard")
bullet("10/10 automated tests confirming all system performance targets met")

doc.add_page_break()

# ── 7. REFERENCES ────────────────────────────────────────────────────────────
heading("7. References")
refs = [
    "Ministry of Railways, Government of India. KAVACH / TCAS Technical Specification. Research Designs and Standards Organisation (RDSO), Lucknow.",
    "Abhishek Pratap Singh. 'Train Collision Avoidance System — KAVACH Training Presentation'. ASTE/P-1/PRYJ, Indian Railways.",
    "Palaksh Bhardwaj. 'Train Collision Avoidance System Based on Generative AI: A Technological Survey'. B.Tech Mini Project, GBU, 2024-25.",
    "Hochreiter, S., Schmidhuber, J. (1997). 'Long Short-Term Memory'. Neural Computation, 9(8), 1735–1780.",
    "Chen, T., Guestrin, C. (2016). 'XGBoost: A Scalable Tree Boosting System'. Proceedings of KDD 2016.",
    "Breiman, L. (2001). 'Random Forests'. Machine Learning, 45(1), 5–32.",
    "FastAPI Documentation. https://fastapi.tiangolo.com",
    "Indian Railways Signal Engineering Manual. Ministry of Railways.",
]
for i, ref in enumerate(refs, 1):
    bullet(f"[{i}] {ref}")

# ── SAVE ─────────────────────────────────────────────────────────────────────
os.makedirs('docs', exist_ok=True)
doc.save('docs/SmartKavach_Final_Report.docx')
print("✅ Final report saved to docs/SmartKavach_Final_Report.docx")